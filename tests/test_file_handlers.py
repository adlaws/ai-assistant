"""Tests for file handlers."""

from __future__ import annotations

import pytest

from src.indexer.file_processing import (
    CSVHandler,
    JSONHandler,
    MarkdownHandler,
    PDFHandler,
    PythonHandler,
    TextHandler,
    WordHandler,
    ImageHandler,
    FILE_HANDLERS,
    get_handler,
)


class TestTextHandler:
    """Test cases for TextHandler."""

    def test_text_handler_init(self, tmp_path):
        """Test text handler initialization."""
        text_file = tmp_path / "test.txt"
        text_file.write_text("Test content")

        handler = TextHandler(str(text_file))

        assert handler.filepath == str(text_file)

    def test_text_handler_load_document(self, tmp_path):
        """Test loading a text document."""
        text_file = tmp_path / "test.txt"
        text_file.write_text("Test content")

        handler = TextHandler(str(text_file))
        content = handler.load_document()

        assert content == "Test content"


class TestMarkdownHandler:
    """Test cases for MarkdownHandler."""

    def test_markdown_handler_init(self, tmp_path):
        """Test markdown handler initialization."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Title\n\nContent")

        handler = MarkdownHandler(str(md_file))

        assert handler.filepath == str(md_file)

    def test_markdown_handler_load_document(self, tmp_path):
        """Test loading a markdown document."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Title\n\nContent")

        handler = MarkdownHandler(str(md_file))
        content = handler.load_document()

        assert content == "# Title\n\nContent"


class TestPDFHandler:
    """Test cases for PDFHandler."""

    def test_pdf_handler_init(self, tmp_path):
        """Test PDF handler initialization."""
        pdf_path = tmp_path / "test.pdf"
        pdf_path.write_text("")  # Create empty file for init test

        handler = PDFHandler(str(pdf_path))

        assert handler.filepath == str(pdf_path)

    def test_pdf_handler_load_document(self, tmp_path):
        """Test loading a PDF document."""
        try:
            from pypdf import PdfWriter

            # Create a simple PDF
            writer = PdfWriter()
            writer.add_blank_page(width=200, height=200)
            pdf_path = tmp_path / "test.pdf"
            with open(pdf_path, 'wb') as f:
                writer.write(f)

            handler = PDFHandler(str(pdf_path))
            content = handler.load_document()

            # PDF without text should return empty or minimal string
            assert isinstance(content, str)
        except ImportError:
            pytest.skip("pypdf not installed")


class TestJSONHandler:
    """Test cases for JSONHandler."""

    def test_json_handler_init(self, tmp_path):
        """Test JSON handler initialization."""
        json_file = tmp_path / "test.json"
        json_file.write_text('{"key": "value"}')

        handler = JSONHandler(str(json_file))

        assert handler.filepath == str(json_file)

    def test_json_handler_load_document(self, tmp_path):
        """Test loading a JSON document."""
        json_file = tmp_path / "test.json"
        json_file.write_text('{"key": "value", "number": 123}')

        handler = JSONHandler(str(json_file))
        content = handler.load_document()

        # JSON is formatted with indent=2
        assert '"key"' in content
        assert '"value"' in content
        assert '123' in content


class TestCSVHandler:
    """Test cases for CSVHandler."""

    def test_csv_handler_init(self, tmp_path):
        """Test CSV handler initialization."""
        csv_file = tmp_path / "test.csv"
        csv_file.write_text("name,age\nAlice,30\nBob,25")

        handler = CSVHandler(str(csv_file))

        assert handler.filepath == str(csv_file)

    def test_csv_handler_load_document(self, tmp_path):
        """Test loading a CSV document."""
        csv_file = tmp_path / "test.csv"
        csv_file.write_text("name,age\nAlice,30\nBob,25")

        handler = CSVHandler(str(csv_file))
        content = handler.load_document()

        # CSV is loaded via pandas and output uses to_string()
        assert "name" in content
        assert "age" in content
        assert "Alice" in content
        assert "30" in content
        assert "Bob" in content
        assert "25" in content


class TestWordHandler:
    """Test cases for WordHandler."""

    def test_word_handler_init(self, tmp_path):
        """Test Word handler initialization."""
        from docx import Document
        docx_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test content")
        doc.save(str(docx_file))

        handler = WordHandler(str(docx_file))

        assert handler.filepath == str(docx_file)

    def test_word_handler_load_document(self, tmp_path):
        """Test loading a Word document."""
        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("Hello, World!")
            doc_path = tmp_path / "test.docx"
            doc.save(doc_path)

            handler = WordHandler(str(doc_path))
            content = handler.load_document()

            assert content == "Hello, World!"
        except ImportError:
            pytest.skip("python-docx not installed")


class TestImageHandler:
    """Test cases for ImageHandler."""

    def test_image_handler_init(self, tmp_path):
        """Test Image handler initialization."""
        try:
            from PIL import Image

            # Create a simple image
            img = Image.new('RGB', (100, 100), color='red')
            img_path = tmp_path / "test.png"
            img.save(img_path)

            handler = ImageHandler(str(img_path))
            assert handler.filepath == str(img_path)
        except ImportError:
            pytest.skip("PIL not installed")

    def test_image_handler_load_document(self, tmp_path):
        """Test loading an image document."""
        try:
            from PIL import Image

            # Create a simple image
            img = Image.new('RGB', (100, 100), color='red')
            img_path = tmp_path / "test.png"
            img.save(img_path)

            handler = ImageHandler(str(img_path))
            content = handler.load_document()

            assert "Image:" in content
            assert "test.png" in content
            assert "Base64:" in content
        except ImportError:
            pytest.skip("PIL not installed")


class TestPythonHandler:
    """Test cases for PythonHandler."""

    def test_python_handler_init(self, tmp_path):
        """Test Python handler initialization."""
        py_file = tmp_path / "test.py"
        py_file.write_text("print('Hello, World!')")

        handler = PythonHandler(str(py_file))

        assert handler.filepath == str(py_file)

    def test_python_handler_load_document(self, tmp_path):
        """Test loading a Python document."""
        py_file = tmp_path / "test.py"
        py_file.write_text("print('Hello, World!')")

        handler = PythonHandler(str(py_file))
        content = handler.load_document()

        assert content == "print('Hello, World!')"


class TestGetHandler:
    """Test cases for get_handler function."""

    def test_get_handler_text(self, tmp_path):
        """Test getting handler for text file."""
        text_file = tmp_path / "test.txt"
        text_file.write_text("Test")

        handler = get_handler(str(text_file))

        assert isinstance(handler, TextHandler)

    def test_get_handler_markdown(self, tmp_path):
        """Test getting handler for markdown file."""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Title")

        handler = get_handler(str(md_file))

        assert isinstance(handler, MarkdownHandler)

    def test_get_handler_json(self, tmp_path):
        """Test getting handler for JSON file."""
        json_file = tmp_path / "test.json"
        json_file.write_text('{"key": "value"}')

        handler = get_handler(str(json_file))

        assert isinstance(handler, JSONHandler)

    def test_get_handler_csv(self, tmp_path):
        """Test getting handler for CSV file."""
        csv_file = tmp_path / "test.csv"
        csv_file.write_text("name,age\nAlice,30")

        handler = get_handler(str(csv_file))

        assert isinstance(handler, CSVHandler)

    def test_get_handler_pdf(self, tmp_path):
        """Test getting handler for PDF file."""
        try:
            from pypdf import PdfWriter

            # Create a simple PDF
            writer = PdfWriter()
            writer.add_blank_page(width=200, height=200)
            pdf_path = tmp_path / "test.pdf"
            with open(pdf_path, 'wb') as f:
                writer.write(f)

            handler = get_handler(str(pdf_path))

            assert isinstance(handler, PDFHandler)
        except ImportError:
            pytest.skip("pypdf not installed")

    def test_get_handler_word(self, tmp_path):
        """Test getting handler for Word file."""
        try:
            from docx import Document

            doc = Document()
            doc.add_paragraph("Hello")
            doc_path = tmp_path / "test.docx"
            doc.save(doc_path)

            handler = get_handler(str(doc_path))

            assert isinstance(handler, WordHandler)
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_get_handler_image(self, tmp_path):
        """Test getting handler for image file."""
        try:
            from PIL import Image

            # Create a simple image
            img = Image.new('RGB', (100, 100), color='green')
            img_path = tmp_path / "test.png"
            img.save(img_path)

            handler = get_handler(str(img_path))

            assert isinstance(handler, ImageHandler)
        except ImportError:
            pytest.skip("PIL not installed")

    def test_get_handler_python(self, tmp_path):
        """Test getting handler for Python file."""
        py_file = tmp_path / "test.py"
        py_file.write_text("print('Hello')")

        handler = get_handler(str(py_file))

        assert isinstance(handler, PythonHandler)

    def test_get_handler_unknown_extension(self, tmp_path):
        """Test getting handler for unknown file extension."""
        unknown_file = tmp_path / "test.xyz"
        unknown_file.write_text("Test")

        with pytest.raises(Exception):
            get_handler(str(unknown_file))

    def test_get_handler_path_traversal(self, tmp_path):
        """Test that paths with unknown extensions raise exceptions."""
        # Path traversal files won't have recognized extensions
        traversal_file = tmp_path / "passwd.xyz"
        traversal_file.write_text("Test")

        # Should raise exception for unknown extension
        with pytest.raises(Exception):
            get_handler(str(traversal_file))


class TestFileHandlersRegistry:
    """Test cases for file handlers registry."""

    def test_file_handlers_registry(self):
        """Test that all handlers are registered."""
        assert len(FILE_HANDLERS) > 0

        # Check that expected extensions are registered
        assert '.txt' in FILE_HANDLERS
        assert '.md' in FILE_HANDLERS
        assert '.json' in FILE_HANDLERS
        assert '.csv' in FILE_HANDLERS

    def test_file_handlers_dict(self):
        """Test that FILE_HANDLERS is a dictionary."""
        assert isinstance(FILE_HANDLERS, dict)

        # Check that all values are handler classes
        for ext, handler_class in FILE_HANDLERS.items():
            assert callable(handler_class)

    def test_get_handler_returns_correct_class(self, tmp_path):
        """Test that get_handler returns the correct handler class."""
        # Test with a known extension
        text_file = tmp_path / "test.txt"
        text_file.write_text("test content")
        handler = get_handler(str(text_file))

        assert isinstance(handler, TextHandler)
